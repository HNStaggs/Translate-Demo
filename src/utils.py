from docx import Document


def extract_text_from_docx(file):

    """Extract text from DOCX file."""

    doc = Document(file)

    full_text = []

    for paragraph in doc.paragraphs:

        if paragraph.text.strip():  # Only add non-empty paragraphs

            full_text.append(paragraph.text.strip())

    return full_text


def chunk_text(text_list, max_length=500):

    """Break text into smaller chunks while preserving paragraph boundaries."""

    chunks = []

    current_chunk = []

    current_length = 0

    

    for paragraph in text_list:

        # If adding this paragraph would exceed max_length, save current chunk and start new one

        if current_length + len(paragraph.split()) > max_length and current_chunk:

            chunks.append(' '.join(current_chunk))

            current_chunk = []

            current_length = 0

        

        current_chunk.append(paragraph)

        current_length += len(paragraph.split())

    

    # Add the last chunk if it exists

    if current_chunk:

        chunks.append(' '.join(current_chunk))

    

    return chunks


def create_translated_document(original_texts, translated_texts):

    """Create a new document with original and translated text."""

    doc = Document()

    

    # Add title

    doc.add_heading('Translation Document', 0)

    

    # Add original and translated text in parallel

    for orig, trans in zip(original_texts, translated_texts):

        doc.add_paragraph("Original Text:")

        doc.add_paragraph(orig)

        doc.add_paragraph("\nTranslated Text:")

        doc.add_paragraph(trans)

        doc.add_paragraph("\n" + "-"*50 + "\n")  # Separator

    

    return doc
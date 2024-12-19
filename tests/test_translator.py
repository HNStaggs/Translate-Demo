import pytest
from src.translator import Translator
from src.utils import chunk_text, extract_text_from_docx
from docx import Document
import io

def test_chunk_text():
    """Test text chunking functionality."""
    text_list = ["This is a test.", "This is another test.", "And one more test."]
    chunks = chunk_text(text_list, max_length=2)
    assert len(chunks) > 1

def test_translator_initialization():
    """Test translator class initialization."""
    translator = Translator()
    assert isinstance(translator.LANGUAGE_MODELS, dict)
    assert len(translator.LANGUAGE_MODELS) > 0

def test_model_loading():
    """Test model loading functionality."""
    translator = Translator()
    model_name = translator.LANGUAGE_MODELS["Spanish"]
    model, tokenizer = translator.load_model(model_name)
    assert model is not None
    assert tokenizer is not None

def create_test_docx():
    """Create a test DOCX file."""
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has multiple paragraphs.")
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def test_docx_extraction():
    """Test DOCX text extraction."""
    doc_bytes = create_test_docx()
    text_list = extract_text_from_docx(doc_bytes)
    assert len(text_list) == 2
    assert "test document" in text_list[0]